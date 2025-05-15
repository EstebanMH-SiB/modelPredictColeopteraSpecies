# -*- coding: utf-8 -*-
"""EMH_Predicción_número_especies_coleoptera_antioquia.py
# Modelo de predicción para el número de especies de Coleoptera en Colombia y en el Departamento de Antioquia


Autor: Esteban Marentes Herrera
Enlace GitHub https://github.com/EstebanMH-SiB/modelPredictColeopteraSpecies
"""

# Importa las librerías necesarias
import os 
import pandas as pd 
import numpy as np
import matplotlib.pyplot as plt
import geopandas as gpd
import tensorflow as tf
import keras
import keras_tuner as kt
import rasterio
from shapely.geometry import Point
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler, OneHotEncoder, LabelEncoder, MinMaxScaler
from sklearn.linear_model import LogisticRegression, LinearRegression
from sklearn.model_selection import train_test_split, cross_val_score, RandomizedSearchCV, KFold, GridSearchCV
from sklearn.decomposition import PCA
from sklearn.metrics import accuracy_score, mean_absolute_error, classification_report, ndcg_score, mean_squared_error, r2_score, precision_score, recall_score, f1_score, confusion_matrix
from sklearn.ensemble import RandomForestRegressor
from sklearn.neural_network import MLPRegressor
from sklearn.svm import SVR
from sklearn.tree import plot_tree
from xgboost import XGBRegressor
from tensorflow.keras import layers, models
from tensorflow.keras.models import Sequential, model_from_json
from tensorflow.keras.layers import Dense, Dropout, Input, Normalization, Rescaling
from tensorflow.keras.callbacks import EarlyStopping
from tensorflow.keras import backend as K
from keras.utils import plot_model
from keras.optimizers import Adam
import seaborn as sns
from scipy.stats import chi2_contingency
from docx import Document
from io import StringIO
from joblib import dump, load


# Definir el directorio de trabajo (modificar según sea necesario)
os.chdir("/Users/estebanmarentes/Desktop/EstebanMH/GBIFColombiaCompleto20200825/MaestriaUJaveriana/TercerSemestre/TesisFinal/Datos")

# Se cargan los datos de coleoptera del mundo en chunks de 100000 datos
chunks_list = []
chunk_size = 1000000

# Iterar sobre los chunks
for chunk in pd.read_csv('coleoptera_completa_gbif.csv',  encoding = "utf8", sep="\t", chunksize=chunk_size, on_bad_lines='skip'):
    chunks_list.append(chunk)
coleoptera_completa_mundo = pd.concat(chunks_list, ignore_index=True) # concatenar los chunks juntos

# Quitar las columnas que no tienen nada que ver con lo biológico
coleoptera_completa_mundo_especie = coleoptera_completa_mundo.drop(columns=['datasetKey', 'publishingOrgKey', 'day', 'month', 'year', 'taxonKey', 'institutionCode', 'catalogNumber', 'recordNumber', 'rightsHolder', 'typeStatus', 'establishmentMeans', 'lastInterpreted', 'mediaType'])

# Quitar los datos sin coordenadas
coleoptera_completa_mundo_especie = coleoptera_completa_mundo_especie.dropna(subset=['decimalLatitude', 'decimalLongitude'])


# Crear la columna tipo punto usando geopandas para hacer el cruce con las capas vectoriales
coleoptera_completa_mundo_especie[['decimalLatitude', 'decimalLongitude']] = coleoptera_completa_mundo_especie[['decimalLatitude','decimalLongitude']].fillna(value=0)
coleoptera_completa_mundo_especie['Coordinates'] = list(zip(coleoptera_completa_mundo_especie.decimalLongitude, coleoptera_completa_mundo_especie.decimalLatitude))
coleoptera_completa_mundo_especie['Coordinates'] = coleoptera_completa_mundo_especie['Coordinates'].apply(Point)
coleoptera_completa_mundo_especie = gpd.GeoDataFrame(coleoptera_completa_mundo_especie, geometry='Coordinates')
coleoptera_completa_mundo_especie.crs = {'init' :'epsg:4326'}


#Realizar el curce con la capa de Colombia
departamento = gpd.read_file("departamento/MGN_DPTO_POLITICO.shp", encoding = "utf8") # cargar el archivo shapefile de los municipios de Colombia
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie, departamento, how="inner", op="intersects")

#Quitar las columnas extras creadas con el cruce y cambiar el nombre por departamento
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'DPTO_CCDGO', 'DPTO_NANO_', 'DPTO_CACTO', 'DPTO_NANO','Shape_Leng', 'Shape_Area', 'DPTO_CNMBR'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'stateProv': 'departamento'}, inplace=True)

# Eliminar las columnas completamente vacías
coleoptera_completa_mundo_especie_colombia.dropna(axis=1, how='all', inplace=True)

#### Cargar los archivos con información adicional
etiquetas_Coleoptera_Colombia = pd.read_csv('EtiquetasColeopteraColombia.txt', encoding = "utf8", sep="\t")
plantae_colombia = pd.read_csv('PlantaeColombiaCompleto/verbatim.txt', encoding = "utf8", sep="\t", usecols=['gbifID', 'occurrenceID',  'eventDate',  'country', 'stateProvince', 'county', 'municipality', 'locality','decimalLatitude', 'decimalLongitude', 'geodeticDatum', 'coordinateUncertaintyInMeters', 'taxonID', 'scientificName', 'genus', 'specificEpithet', 'infraspecificEpithet', 'taxonRank', 'scientificNameAuthorship'])
Radiacion_solar_global_promedio_multianual = gpd.read_file("Radiacion_solar_global_promedio_multianual/Shape/GDBIDEAM.CN2_Rdcn_Solar_Global_ProAnual.shp", encoding = "utf8")
Humedad_Relativa_Anual_Promedio_Multianual_1981_2010 = gpd.read_file("Humedad_Relativa_Anual_Promedio_Multianual_1981_2010/SHP/ACC2014_HmRl_AA_MA_1981_2010.shp", encoding = "utf8")
Temperatura_Maxima_Media_Mensual_Promedio_Multianual_1981_2010 = gpd.read_file("Temperatura_Maxima_Media_Mensual_Promedio_Multianual_1981_2010/SHP/ACC2014_TmMx_MM_1981_2010_01.shp", encoding = "utf8")
Temperatura_Media_Mensual_Promedio_Multianual_1981_2010 = gpd.read_file("Temperatura_Media_Mensual_Promedio_Multianual_1981_2010/SHP/ACC2014_TmMd_MM_1981_2010_01.shp", encoding = "utf8")
Temperatura_Minima_Media_Mensual_Promedio_Multianual_1981_2010 = gpd.read_file("Temperatura_Minima_Media_Mensual_Promedio_Multianual_1981_2010/SHP/ACC2014_TmMn_MM_1981_2010_01.shp", encoding = "utf8")
Velocidad_viento_10_mtrs_altura_Mensual_2000_2010 = gpd.read_file("Velocidad_viento_10_mtrs_altura_Mensual_2000_2010/SHP/AVC2014_VlVn_10m_MM_2000_2010_01.shp", encoding = "utf8")
ECCMC_Ver21_100K = gpd.read_file("Mapa_ecosistemas_Continentales_Marinos_Costeros_100K_V2.1_2017/Shape_E_ECCMC_Ver21_100K/E_ECCMC_Ver21_100K.shp", encoding = "utf8")

### Realizar los cruces con las distintas capas

#Realizar cruces Geo radiacion
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, Radiacion_solar_global_promedio_multianual, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'OBJECTID', 'ID', 'GRIDCODE', 'Shape_Leng', 'Shape_Area','RULEID'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'RANGO': 'Radiacion_solar_global_promedio_multianual'}, inplace=True)

#Realizar cruces Geo Humedad relativa
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, Humedad_Relativa_Anual_Promedio_Multianual_1981_2010, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'OBJECTID', 'GRIDCODE', 'Shape_Leng', 'Shape_Area','RULEID'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'RANGO': 'Humedad_Relativa_Anual_Promedio_Multianual_1981_2010'}, inplace=True)

#Realizar cruces Geo Temperatura maxima
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, Temperatura_Maxima_Media_Mensual_Promedio_Multianual_1981_2010, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'PeriodoIni', 'PeriodoFin'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'RANGO': 'Temperatura_Maxima_Media_Mensual_Promedio_Multianual_1981_2010'}, inplace=True)

#Realizar cruces Geo Temperatura media
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, Temperatura_Media_Mensual_Promedio_Multianual_1981_2010, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'GRIDCODE_right','PeriodoIni', 'PeriodoFin'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'RANGO': 'Temperatura_Media_Mensual_Promedio_Multianual_1981_2010'}, inplace=True)

#Realizar cruces Geo Temperatura mínima
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, Temperatura_Minima_Media_Mensual_Promedio_Multianual_1981_2010, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'GRIDCODE','PeriodoIni', 'PeriodoFin'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'RANGO': 'Temperatura_Minima_Media_Mensual_Promedio_Multianual_1981_2010'}, inplace=True)

#Realizar cruces Velocidad viento
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, Velocidad_viento_10_mtrs_altura_Mensual_2000_2010, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'GRIDCODE','GRIDCODE_left','PeriodoIni', 'PeriodoFin'])
coleoptera_completa_mundo_especie_colombia.rename(columns={'RANGO': 'Velocidad_viento_10_mtrs_altura_Mensual_2000_2010'}, inplace=True)

#Realizar cruces Ecosistemas Continentales
coleoptera_completa_mundo_especie_colombia = gpd.sjoin(coleoptera_completa_mundo_especie_colombia, ECCMC_Ver21_100K, how="left", op="intersects")
#Quitar las columnas extras creadas con el cruce y cambiar el nombre de la columna con información
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['index_right', 'OBJECTID','TIPO_ECOSI','GRADO_TRAN', 'GRAN_BIOMA', 'BIOMA_PREL','BIOMA_IAvH','ECOS_SINTE', 'AMBIENTE_A','SUBSISTEMA','ZONA_HIDRO', 'ORIGEN', 'TIPO_AGUA','PAISAJE','AMB_EDAFOG', 'DESC_AMB_E', 'SUSTRATO', 'ZONA', 'TEMPERATUR', 'SALINIDAD', 'PROVINCIA', 'ECO_REGION', 'ECO_ZONA', 'ORIGEN_MAR', 'CONFIGURAC', 'CLAS_BIOTI', 'SUBCLAS_BI', 'GRUPO_BIOT', 'SECTORES', 'Area_ha', 'UNI_BIOTIC', 'ANFIBIOS', 'AVES', 'MAGNOLIOPS', 'MAMIFEROS', 'REPTILES', 'SHAPE_Leng', 'SHAPE_Area', 'RULEID'])

#Realizar cruces DEM elevación
ruta_dem = r"/Users/estebanmarentes/Desktop/EstebanMH/GBIFColombiaCompleto20200825/MaestriaUJaveriana/TercerSemestre/TesisFinal/Datos/astergdem2.tif"
with rasterio.open(ruta_dem) as src:
    coords = [(x, y) for x, y in zip(coleoptera_completa_mundo_especie_colombia.geometry.x, coleoptera_completa_mundo_especie_colombia.geometry.y)]
    valores = [v[0] for v in src.sample(coords, indexes=1)]
    
coleoptera_completa_mundo_especie_colombia["elevacion_dem"] = valores

#Realizar cruces raster de Precipitación
ruta_precipitacion = r"/Users/estebanmarentes/Desktop/EstebanMH/GBIFColombiaCompleto20200825/MaestriaUJaveriana/TercerSemestre/TesisFinal/Datos/Escenario_Precipitacion_1976_2005/ECC_Prcp_GeoTiff_2011_2040/ECC_Prcp_1976_2005_100K_2015.tif"
with rasterio.open(ruta_precipitacion) as src:
    coords = [(x, y) for x, y in zip(coleoptera_completa_mundo_especie_colombia.geometry.x, coleoptera_completa_mundo_especie_colombia.geometry.y)]
    valores = [v[0] for v in src.sample(coords, indexes=1)]

coleoptera_completa_mundo_especie_colombia["ECC_Prcp_1976_2005_100K_2015"] = valores



### Cargar los datos de plantas y limpiarlos para dejar solamente las especies, los que tengan coordenadas y dejar datos sin duplicados
plantae_colombia = pd.read_csv('PlantaeColombiaCompleto/verbatim.txt', encoding = "utf8", sep="\t", usecols=['gbifID', 'occurrenceID',  'decimalLatitude', 'decimalLongitude','scientificName', 'taxonRank'])
plantae_colombia['taxonRank']=plantae_colombia['taxonRank'].str.title()
plantae_colombia['taxonRank']=plantae_colombia['taxonRank'].replace('Specie', 'Especie').replace('Espécie', 'Especie').replace('Epecie', 'Especie').replace('Especies', 'Especie').replace('Species', 'Especie')
plantae_colombia_especie=plantae_colombia.loc[plantae_colombia['taxonRank'] == 'Especie' ]
#Dejar solo las filas con información
plantae_colombia_especie = plantae_colombia_especie.dropna(subset=['decimalLatitude', 'decimalLongitude', 'scientificName'])
# Borrar las filas duplicadas
plantae_colombia_especie = plantae_colombia_especie.drop(columns=['occurrenceID'])
plantae_colombia_especie = plantae_colombia_especie.drop_duplicates(subset=[col for col in plantae_colombia_especie.columns if col != 'gbifID'])
# Limpiar las coordenadas
plantae_colombia_especie = plantae_colombia_especie[~plantae_colombia_especie['decimalLongitude'].str.contains('-', na=False)]
plantae_colombia_especie = plantae_colombia_especie[~plantae_colombia_especie['decimalLatitude'].isin(['71.122.222', '73.275.528','4.525694.'])]
plantae_colombia_especie[['decimalLongitude']]=plantae_colombia_especie[['decimalLongitude']].replace(',', '.', regex=True).replace(' ', '', regex=True).astype(float)
plantae_colombia_especie[['decimalLatitude']]=plantae_colombia_especie[['decimalLatitude']].replace(',', '.', regex=True).replace(' ', '', regex=True).astype(float)
plantae_colombia_especie['decimalLongitude'] = pd.to_numeric(plantae_colombia_especie['decimalLongitude'], errors='coerce')
plantae_colombia_especie['decimalLatitude'] = pd.to_numeric(plantae_colombia_especie['decimalLatitude'], errors='coerce')
# Trasnformar a GPD
plantae_colombia_especie[['decimalLatitude', 'decimalLongitude']] = plantae_colombia_especie[['decimalLatitude','decimalLongitude']].fillna(value=0)
plantae_colombia_especie['Coordinates'] = list(zip(plantae_colombia_especie.decimalLongitude, plantae_colombia_especie.decimalLatitude))
plantae_colombia_especie['Coordinates'] = plantae_colombia_especie['Coordinates'].apply(Point)
plantae_colombia_especie = gpd.GeoDataFrame(plantae_colombia_especie, geometry='Coordinates')
plantae_colombia_especie.crs = {'init' :'epsg:4326'}
# Cruzar a Colombia
plantae_colombia_especie = gpd.sjoin(plantae_colombia_especie, departamento, how="inner", op="intersects")
plantae_colombia_especie = plantae_colombia_especie.drop(columns=['index_right', 'DPTO_CCDGO', 'DPTO_NANO_', 'DPTO_CACTO', 'DPTO_NANO','Shape_Leng', 'Shape_Area', 'DPTO_CNMBR'])


# Dejar solamente las columnas de coordenadaas y gbifID para crear el buffer
gdf = coleoptera_completa_mundo_especie_colombia[[ "Coordinates", "gbifID"]]
#Crear un buffer aproximado de 1000 metros de los puntos de Coleoptera
gdf['buffer'] = gdf.geometry.buffer(0.01)
gdf = gpd.GeoDataFrame(gdf, geometry='buffer')
# Dejar solamente las columnas de nombre científico y gbifID para los puntos de plantas
plantae_cruce = plantae_colombia_especie[[ "gbifID", "scientificName", "Coordinates"]]
plantae_cruce.rename(columns={'scientificName': 'especiePlanta'}, inplace=True) #cambiar nombre a especies plantas
# Cruce de los datos entre buffers y los puntos
coleopteros_colombia_plantas = gpd.sjoin_nearest(gdf, plantae_cruce , how='left')
# Crear la función para concatenar los datos
def concatenate_values(group):
    return pd.Series({
        'column1_concat': ' | '.join(group['especiePlanta'].astype(str))
    })
# Agrupar las especies por 'gbifID' y aplicar la función de concatenación
result = coleopteros_colombia_plantas.groupby('gbifID_left').apply(concatenate_values).reset_index()
# Cambiar los nombres de las columnas
result.rename(columns={'column1_concat': 'especiePlanta'}, inplace=True)
result.rename(columns={'gbifID_left': 'gbifID'}, inplace=True)

# Juntar la columna creada con el dataset Completo
coleoptera_completa_mundo_especie_colombia = pd.merge(coleoptera_completa_mundo_especie_colombia, result,  on='gbifID',  how='left')



# Exportar los datos completos
coleoptera_completa_colombia = coleoptera_completa_mundo_especie_colombia
coleoptera_completa_colombia.to_csv( 'coleoptera_completa'+'.txt', sep="\t", encoding = "utf8", index=False)

#### Quitar los NA
numeric_imputer = SimpleImputer(strategy='median') # Mediana para los datos numéricos
categorical_imputer = SimpleImputer(strategy='most_frequent') # Moda para los datos categóricos

# Hacer un loop para cada columnas y aplicar la imputación adecuada
for column in coleoptera_completa_mundo_especie_colombia.columns:
    if coleoptera_completa_mundo_especie_colombia[column].dtype in ['float64', 'int64']:  # Numerica
        coleoptera_completa_mundo_especie_colombia[column] = numeric_imputer.fit_transform(coleoptera_completa_mundo_especie_colombia[[column]]).flatten() 
    else:  # Categóricas
        coleoptera_completa_mundo_especie_colombia[column] = categorical_imputer.fit_transform(coleoptera_completa_mundo_especie_colombia[[column]]).flatten()

# Exportar los datos sin NA
coleoptera_completa_colombia = coleoptera_completa_mundo_especie_colombia
coleoptera_completa_colombia.to_csv( 'coleoptera_completa_sinna'+'.txt', sep="\t", encoding = "utf8", index=False)

# Etiquetar con los datos con la información de las familias
etiquetas_Coleoptera_Colombia = etiquetas_Coleoptera_Colombia[['family','NumeroEspeciesFamilia']] #Quedarse solo con la columna del número de registros
coleoptera_completa_mundo_especie_colombia = pd.merge(coleoptera_completa_mundo_especie_colombia,etiquetas_Coleoptera_Colombia,how='left',on='family')

#### Crear un transformador para codificar las características categóricas

# Seleccionar solamente las variables categóricas
categorical_columns = coleoptera_completa_mundo_especie_colombia.select_dtypes(include=['object']).columns
# Seleccionar solamente variables las numéricas
numerical_columns = coleoptera_completa_mundo_especie_colombia.select_dtypes(include=['float64']).columns

label_encoders = {}

for column in categorical_columns:
    le = LabelEncoder()  # Crear un LabelEncoder para cada columnas
    coleoptera_completa_mundo_especie_colombia[column] = coleoptera_completa_mundo_especie_colombia[column].astype(str) # convertir los valores faltantes a string
    coleoptera_completa_mundo_especie_colombia[column] = le.fit_transform(coleoptera_completa_mundo_especie_colombia[column])
    label_encoders[column] = le # Guardar los diccionarios para poder hacer la trasnformación inversa luego
    
# Exportar los datos codificados con LE
coleoptera_completa_colombia = coleoptera_completa_mundo_especie_colombia
coleoptera_completa_colombia.to_csv( 'coleoptera_completa_le'+'.txt', sep="\t", encoding = "utf8")


#### descartar los datos no etiquetados y borrar las columnas completamente vacias del nuevo dataset#######
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.dropna(subset=['NumeroEspeciesFamilia'])
coleoptera_completa_mundo_especie_colombia.dropna(axis=1, how='all', inplace=True)


# Exportar los datos ajustados con etiquetas
coleoptera_completa_mundo_especie_colombia.to_csv( 'coleoptera_completa_mundo_especie_colombia_etiquetados'+'.txt', sep="\t", encoding = "utf8")


####################################################################################################################


### hacer un plot del mapa de Colombia par ver los datos
departamento =gpd.read_file("departamento/MGN_DPTO_POLITICO.shp", encoding = "utf8") # cargar el archivo shapefile de los municipios de Colombia
# corret todo el bloque junto
fig, ax = plt.subplots(figsize=(10, 10))
departamento.plot(ax=ax, color='lightblue', edgecolor='black', label="Departamento")
coleoptera_completa_mundo_especie_colombia.plot(ax=ax, color='red', marker='o', markersize=5, label="Registros de especies de Coleoptera")
ax.set_title("Mapa de Colombia con registros de Coleoptera")
ax.set_xlabel("Longitud")
ax.set_ylabel("Latitud")
ax.legend()
plt.savefig('mapaColeoptera.png', dpi=300, bbox_inches='tight')
plt.show()


## Vista final de los datos y exportar una muestra
nan_count = coleoptera_completa_mundo_especie_colombia.isna().sum(axis=1)
vista_final = coleoptera_completa_mundo_especie_colombia[nan_count <= 4]
vista_final = vista_final.head(10)
vista_final = vista_final.transpose()
vista_final.to_excel("dataframe.xlsx", index=True)


### generar descriptivos del conjunto de datos

## Etapa 2: Análisis exploratorio de los datos.

# Información general sobre el dataset
print("Información del Dataset:")
print(coleoptera_completa_mundo_especie_colombia.info())

# Resumen estadístico de las características numéricas
print("\nResumen Estadístico de las Características Numéricas:")
print(coleoptera_completa_mundo_especie_colombia.describe())

# Resumen de las características categóricas
print("\nResumen de las Características Categóricas:")
print(coleoptera_completa_mundo_especie_colombia.describe(include=['object']))


def describe_dataframe_with_categorical_to_word(df, filename='descripcion_df.docx'):
    pd.set_option('display.max_columns', None)
    pd.set_option('display.width', None)
    
    # Crear documento
    doc = Document()
    
    # Agregar título
    doc.add_heading('Descripción general del conjunto de datos', 0)
    
    # Agregar información básica
    doc.add_heading('1. Información básica (Shape, Columns, Data Types):', level=1)
    doc.add_paragraph(f"Shape: {df.shape}")
    doc.add_paragraph(f"Columns: {list(df.columns)}")
    doc.add_paragraph(f"Data Types:\n{df.dtypes}")
    
    # Agregar estadísticas de resumen variables numéricas
    doc.add_heading('2. Estadísticas de resumen variables numéricas:', level=1)
    numeric_summary = df.describe(include=[float, int])  # Show all numeric columns
    doc.add_paragraph(str(numeric_summary))
    
    # Agregar estadísticas de resumen variables categóricas
    doc.add_heading('3. Estadísticas de resumen variables categóricas:', level=1)
    categorical_summary = df.describe(include=[object])  # Show all categorical columns
    doc.add_paragraph(str(categorical_summary))
    
    # Información valores faltantes
    doc.add_heading('4. Información valores faltantes:', level=1)
    missing_values = df.isnull().sum()
    doc.add_paragraph(f"Missing Values per Column:\n{missing_values}")
    
    # Información registros y uso de memoria
    doc.add_heading('5. Información registros y uso de memoria:', level=1)
    buffer = StringIO()
    df.info(buf=buffer)
    doc.add_paragraph(buffer.getvalue())
    
    # Conteo de valores únicos por colúmnas
    doc.add_heading('6. Conteo de valores únicos por colúmnas:', level=1)
    unique_values_count = df.nunique()
    doc.add_paragraph(f"{unique_values_count}")
    
    # Agregar las primeras 5 filas
    doc.add_heading('7. Primeras 5 filas:', level=1)
    doc.add_paragraph(f"{df.head()}")
    
    # Salvar el documento
    doc.save(filename)
    print(f"Descripción guardada a {filename}")

# Exportar los datos a un word
describe_dataframe_with_categorical_to_word(coleoptera_completa_mundo_especie_colombia)


# Seleccionar solamente las variables categóricas
categorical_columns = coleoptera_completa_mundo_especie_colombia.select_dtypes(include=['object']).columns
# Seleccionar solamente variables las numéricas
numerical_columns = coleoptera_completa_mundo_especie_colombia.select_dtypes(include=['float64']).columns


# Crea gráficos de distribución para cada variable numérica
plt.figure(figsize=(18, 6))  # Ajustar el tamaño de la figura para múltiples subplots

for i, feature in enumerate(numerical_columns, 1):
    plt.subplot(1, len(numerical_columns), i)
    sns.histplot(coleoptera_completa_mundo_especie_colombia[feature], kde=True, bins=30)  # kde=True añade la curva de densidad
    plt.title(f'Distribución de {feature}')
    plt.xlabel(feature)
    plt.ylabel('Frecuencia')

plt.tight_layout()  # Ajustar el espaciado entre subplots
plt.show()



def cramers_v(confusion_matrix):
    chi2 = chi2_contingency(confusion_matrix)[0]
    n = confusion_matrix.sum()
    phi2 = chi2 / n
    r, k = confusion_matrix.shape
    phi2corr = max(0, phi2 - ((k - 1)*(r - 1)) / (n - 1))
    rcorr = r - ((r - 1)**2) / (n - 1)
    kcorr = k - ((k - 1)**2) / (n - 1)
    return np.sqrt(phi2corr / min((kcorr - 1), (rcorr - 1)))

categorical_corr = pd.DataFrame(np.zeros((len(categorical_columns), len(categorical_columns))), index=categorical_columns, columns=categorical_columns)

for col1 in categorical_columns:
    for col2 in categorical_columns:
        if col1 == col2:
            categorical_corr.loc[col1, col2] = 1.0
        else:
            confusion_matrix = pd.crosstab(coleoptera_completa_mundo_especie_colombia[col1], coleoptera_completa_mundo_especie_colombia[col2])
            categorical_corr.loc[col1, col2] = cramers_v(confusion_matrix.to_numpy())

# Grafica la matriz de Cramér's V
plt.figure(figsize=(12, 10))
sns.heatmap(categorical_corr, annot=True, cmap='coolwarm', vmin=0, vmax=1)
plt.title('Matriz de Asociación de Cramér\'s V para Variables Categóricas')
plt.show()




## Vista final de los datos y exportar una muestra luego de la codificación
vista_final_le = coleoptera_completa_mundo_especie_colombia.head(5)
vista_final_le = vista_final_le.transpose()
vista_final_le.to_excel("dataframe_le.xlsx", index=True)



# Se eliminan las columnas que se identificaron como redundantes en el conjunto de datos
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['Coordinates','gbifID','occurrenceID','kingdom','phylum','class', 'order', 'stateProvince','infraspecificEpithet','license'])
coleoptera_completa_colombia = coleoptera_completa_colombia.drop(columns=['Coordinates','gbifID','occurrenceID','kingdom','phylum','class', 'order', 'stateProvince','infraspecificEpithet','license'])

# Exportar los datos finales
coleoptera_completa_colombia.to_csv( 'coleoptera_completa_final'+'.txt', sep="\t", encoding = "utf8")
coleoptera_completa_mundo_especie_colombia.to_csv( 'coleoptera_completa_mundo_especie_colombia_etiquetados_final'+'.txt', sep="\t", encoding = "utf8")


####  Se divide en los conjuntos de entrenamiento y prueba.
coleoptera_completa_mundo_especie_colombia = pd.read_csv('coleoptera_completa_mundo_especie_colombia_etiquetados_final.txt', encoding = "utf8", sep="\t")
coleoptera_completa_mundo_especie_colombia = coleoptera_completa_mundo_especie_colombia.drop(columns=['Unnamed: 0'])
# Dejar solamente los datos etiquetados en Y
y = coleoptera_completa_mundo_especie_colombia['NumeroEspeciesFamilia'].values 
arg = tf.convert_to_tensor(y, dtype=tf.float32)
y = y.astype(np.float32) 

# Dejar en X todos los datos menos la columna con las etiquetas
X = coleoptera_completa_mundo_especie_colombia.loc[:, coleoptera_completa_mundo_especie_colombia.columns != "NumeroEspeciesFamilia"] 

# Primer split de 75%
X_train, X_temp, y_train, y_temp = train_test_split(X, y, test_size=0.25, random_state=42)
# Segundo split del 25% restante en test y validación
X_val, X_test, y_val, y_test = train_test_split(X_temp, y_temp, test_size=0.8, random_state=42)



## Etapa 3: Entrenamiento de los modelos

#### Modelo número 0, regresión linear múltiple como base para las comparaciones

model_linear = LinearRegression()
model_linear.fit(X_train, y_train)
y_pred = model_linear.predict(X_test)

# Evaluar el modelo
mse = mean_squared_error(y_test, y_pred)
r2 = r2_score(y_test, y_pred)
mae = mean_absolute_error(y_test, y_pred)

print("Mean Absolute Error:", mae)
print("Mean Squared Error:", mse)
print("R-squared:", r2)

# Exportar modelo lineal
dump(model_linear, 'modelo_lineal_entrenado.pkl')
# model_linear = load('modelo_lineal_entrenado.pkl') # usar para cargar el modelo


#### Modelo número 1, perceptrón multicapa
model = Sequential()
model.add(Input(shape=(48,)))  
model.add(Dense(64, activation='relu'))
model.add(Dense(1))


model.summary()
optimizer = Adam(learning_rate=1e-4)
model.compile(optimizer=optimizer, loss='mse', metrics=['mae'])

history = model.fit(X_train, y_train, epochs=50, validation_data=(X_test, y_test))
y_pred = model.predict(X_test)

# Evaluar el modelo
loss, mae = model.evaluate(X_test, y_test)
r2 = r2_score(y_test, y_pred)
mse = mean_squared_error(y_test, y_pred)
print(f"Test Loss: {loss}, Test MAE: {mae}")
print("R-squared:", r2)
print("Mean Squared Error:", mse)

# Exportar el modelo DNN sencillo
model.save('dnn_simple.keras')
dnn_simple = tf.keras.models.load_model('dnn_simple.keras')

plot_model(model, to_file='dnn_simple.png', show_shapes=True, show_layer_names=True)

#### Modelo número 2, red neuronal profunda sin datos estandarizados optimizando los hiperparámetros
np.random.seed(42)
K.clear_session()

# Definir la función de KerasTuner para optimizar los hiperparámetros
def build_model(hp):
    model = Sequential()

    # Primera capa oculta con número de neuronas y función de activación nuneable
    model.add(Dense(units=hp.Int('units', min_value=32, max_value=128, step=32),
                    activation=hp.Choice('activation_0', values=['relu', 'tanh', 'elu', 'leaky_relu', 'swish']),
                    input_dim=X_train.shape[1]))

    # Agregar capas ocultas entre 1 a 5
    for i in range(hp.Int('num_layers', 1, 5)):
        model.add(Dense(units=hp.Int(f'units_{i}', min_value=32, max_value=128, step=32),
                        activation=hp.Choice(f'activation_{i}', values=['relu', 'tanh', 'elu', 'leaky_relu', 'swish'])))
        model.add(Dropout(hp.Float(f'dropout_rate_{i}', min_value=0.2, max_value=0.5, step=0.1)))  # Capa dropout tuneable

    # Capa de salida de regresión simple, no es tuneable
    model.add(Dense(1))

    # Compilar el modelo con un learning rate tuneable
    model.compile(optimizer=Adam(learning_rate=hp.Float('learning_rate', min_value=1e-5, max_value=1e-2, sampling='LOG')),
                  loss='mean_squared_error', metrics=['mae'])

    return model

# Creae el tuner de Hiperbanda en KerasTuner
tuner = kt.Hyperband(
    build_model,
    objective='val_loss',  # Objetivo val-loss
    max_epochs=50,          # Máximo número de épocas por cada intento
    hyperband_iterations=3, # Número de búsqueda en las iteraciones
    directory='my_dir',
    project_name='dnn_tuning_sines'
)

# Buscar los mejores hiperparámetros
tuner.search(X_train, y_train, epochs=50, batch_size=32, validation_data=(X_test, y_test))

# Obtener el mejor modelo y los mejores hiperparámetros
best_model = tuner.get_best_models(num_models=1)[0]
best_hyperparameters = tuner.oracle.get_best_trials(num_trials=1)[0].hyperparameters.values

# Imprimir los mejores hiperparámetros
print("Mejores hiperparámetros: ", best_hyperparameters)

# Evaluar el mejor modelo en los datos de test
test_loss, test_mae = best_model.evaluate(X_test, y_test, verbose=0)
print(f"Test Loss: {test_loss}")
print(f"Test MAE: {test_mae}")


# Exportar el modelo DNN sin escalar
best_model.save('best_model_sinescalar.keras')
best_model_dnn = tf.keras.models.load_model('best_model_sinescalar.keras')
# Mostrar la arquitectura del mejor modelo
best_model_dnn.summary()

plot_model(best_model_dnn, to_file='best_model_dnn.png', show_shapes=True, show_layer_names=True)


# Evaluar el modelo
y_pred = best_model_dnn.predict(X_test)
r2 = r2_score(y_test, y_pred)
loss, mae = best_model_dnn.evaluate(X_test, y_test)
# Imprimir los valores de r2 y la función de pérdida
print("R-squared:", r2)
print(f"Test Loss: {loss}, Test MAE: {mae}")


historySinEscalar = best_model_dnn.fit(X_train, y_train, epochs=50, batch_size=12, validation_data=(X_test, y_test))



# Resumen del modelo
best_model.summary()

# Graficar los resultados
plt.plot(historySinEscalar.history['mae'])
plt.plot(historySinEscalar.history['val_mae'])
plt.title('model mae')
plt.ylabel('mae')
plt.xlabel('epoch')
plt.legend(['train', 'test'], loc='upper left')
plt.show()

plt.plot(historySinEscalar.history['loss'])
plt.plot(historySinEscalar.history['val_loss'])
plt.title('model loss')
plt.ylabel('loss')
plt.xlabel('epoch')
plt.legend(['train', 'test'], loc='upper left')
plt.show()


##### Modelo número 3, red neuronal profunda con los datos escalados optimizando los hiperparámetros
np.random.seed(42)
K.clear_session()
# Escalamiento de los datos
scaler = StandardScaler()
X_train = scaler.fit_transform(X_train)
X_test = scaler.transform(X_test)

# Definir la función de KerasTuner para optimizar los hiperparámetros
def build_model(hp):
    model = Sequential()

    # Primera capa oculta con número de neuronas y función de activación nuneable
    model.add(Dense(units=hp.Int('units', min_value=32, max_value=128, step=32),
                    activation=hp.Choice('activation_0', values=['relu', 'tanh', 'elu', 'leaky_relu', 'swish']),
                    input_dim=X_train.shape[1]))

    # Agregar capas ocultas entre 1 a 5
    for i in range(hp.Int('num_layers', 1, 5)):
        model.add(Dense(units=hp.Int(f'units_{i}', min_value=32, max_value=128, step=32),
                        activation=hp.Choice(f'activation_{i}', values=['relu', 'tanh', 'elu', 'leaky_relu', 'swish'])))
        model.add(Dropout(hp.Float(f'dropout_rate_{i}', min_value=0.2, max_value=0.5, step=0.1)))  # Capa dropout tuneable

    # Capa de salida de regresión simple, no es tuneable
    model.add(Dense(1)) 

    # Compilar el modelo con un learning rate tuneable
    model.compile(optimizer=Adam(learning_rate=hp.Float('learning_rate', min_value=1e-5, max_value=1e-2, sampling='LOG')),
                  loss='mean_squared_error', metrics=['mae'])

    return model

# Creae el tuner de Hiperbanda en KerasTuner
tuner = kt.Hyperband(
    build_model,
    objective='val_loss',  # Objetivo val-loss
    max_epochs=50,         # Máximo número de épocas por cada intento
    hyperband_iterations=3, # Número de búsqueda en las iteraciones
    directory='my_dir',
    project_name='dnn_escalado_escalada'
)

# Buscar los mejores hiperparámetros
tuner.search(X_train, y_train, epochs=50, batch_size=32, validation_data=(X_test, y_test))

# Obtener el mejor modelo y los mejores hiperparámetros
best_model = tuner.get_best_models(num_models=1)[0]
best_hyperparameters = tuner.oracle.get_best_trials(num_trials=1)[0].hyperparameters.values

# Imprimir los mejores hiperparámetros
print("Mejores hiperparámetros: ", best_hyperparameters)

# Evaluar el mejor modelo en los datos de test
test_loss, test_mae = best_model.evaluate(X_test, y_test, verbose=0)
print(f"Test Loss: {test_loss}")
print(f"Test MAE: {test_mae}")


# Salvar el mejor modelo
best_model.save('best_model_escalado.keras')
best_model_dnn_escalado = tf.keras.models.load_model('best_model_escalado.keras')
model_json = best_model_dnn_escalado.to_json()
with open("model_architecture.json", "w") as json_file:
    json_file.write(model_json)
# Mostrar la arquitectura del mejor modelo
best_model_dnn_escalado.summary()


y_pred = best_model_dnn_escalado.predict(X_test)

yprediccion_especies = best_model_dnn_escalado.predict(coleoptera_completa_colombia)


r2 = r2_score(y_test, y_pred)
loss, mae = best_model_dnn_escalado.evaluate(X_test, y_test)

# Imprimir los valores de r2 y la función de pérdida

print("R-squared:", r2)
print(f"Test Loss: {loss}, Test MAE: {mae}")


historyEscalado = best_model_dnn_escalado.fit(X_train, y_train, epochs=50, batch_size=12, validation_data=(X_test, y_test))



# Resumen del modelo
best_model.summary()

# Graficar los resultados
plt.plot(historyEscalado.history['mae'])
plt.plot(historyEscalado.history['val_mae'])
plt.title('model mae')
plt.ylabel('mae')
plt.xlabel('epoch')
plt.legend(['train', 'test'], loc='upper left')
plt.show()

plt.plot(historyEscalado.history['loss'])
plt.plot(historyEscalado.history['val_loss'])
plt.title('model loss')
plt.ylabel('loss')
plt.xlabel('epoch')
plt.legend(['train', 'test'], loc='upper left')
plt.show()


#### Modelo número 4, Random Forest optimizando los hiperparámetros

# Definir una grilla con la distribución posible de los parámetros a usar
rs_param_grid = {
    "n_estimators": list((range(20, 200))), # número de árboles en el forest
    "max_depth": list((range(3, 12))),# Máxima profundidad del árbol
    "min_samples_split": list((range(2, 5))),  # Mínimo númeor de muestras requeridas para separar un nodo
    "min_samples_leaf": list((range(1, 5))),  # Mínimo númeor de muestras requeridas en una hoja del nodo
    "ccp_alpha": [0, 0.001, 0.01, 0.1],  # Costo de complejidad
}

# Crear el RandomForestRegressor
rf = RandomForestRegressor(random_state=42)

# Inicialiar la búsqueda por grilla RandomizedSearchCV() con el rf y la grilla
rf_rs = RandomizedSearchCV(
    estimator=rf,
    param_distributions=rs_param_grid,
    cv=10,  # Number of folds
    n_iter=10,  # Number of parameter candidate settings to sample
    verbose=2,  # The higher this is, the more messages are outputed
    scoring="neg_mean_absolute_error",  # Metric to evaluate performance
    random_state=42
)

# Train the model on the training set
rf_rs.fit(X_train, y_train)

# Print the best parameters and highest accuracy
print("Best parameters found: ", rf_rs.best_params_)
print("Best performance: ", rf_rs.best_score_)

# Predict the outcomes on the test set
y_pred = rf_rs.predict(X_test)
print("Mean Absolute Error:", mean_absolute_error(y_test, y_pred))
print("Mean Squared Error:", mean_squared_error(y_test, y_pred))
print("Root square:", r2_score(y_test, y_pred))



# Salvar el mejor modelo RF
dump(rf_rs, 'best_model_randomforest_v2.joblib')
best_model_rf = load('best_model_randomforest_v2.joblib')

# Obtener el mejor modelo del RandomizedSearchCV
best_rf = rf_rs.best_estimator_

# Extraer el primer árbol del random forest
tree = best_rf.estimators_[0] 

# Plot it
plt.figure(figsize=(20, 10))  # You can adjust the size
plot_tree(
    tree,
    feature_names=X.columns,   # if X is a DataFrame
    filled=True,
    rounded=True
)
plt.title("Random Forest - Tree 0")
plt.savefig("random_forest_tree_0.png", dpi=600, bbox_inches='tight')
plt.show()


## Error de predicion

plt.scatter(y_test, y_pred, alpha=0.5)
plt.xlabel("Valor real")
plt.ylabel("Predicción")
plt.title("Random Forest: Valor real vs Predicción")
plt.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'k--')
plt.show()


###### Validación cruzada k-fold del mejor modelo DNN escalado #####


# Convertir X, y a arrays de numpy si no lo son
X = np.array(X)
y = np.array(y)

model = tf.keras.models.load_model('best_model_escalado.keras')
kf = KFold(n_splits=10, shuffle=True, random_state=42)

r2_scores = []
losses = []
maes = []

for train_idx, test_idx in kf.split(X):
    X_train, X_test = X[train_idx], X[test_idx]
    y_train, y_test = y[train_idx], y[test_idx]

    # Escalar los datos con parámetros del fold actual
    scaler = StandardScaler()
    X_train = scaler.fit_transform(X_train)
    X_test = scaler.transform(X_test)

    # Predecir y evaluar
    y_pred = model.predict(X_test).flatten()
    r2 = r2_score(y_test, y_pred)
    loss, mae = model.evaluate(X_test, y_test, verbose=0)

    r2_scores.append(r2)
    losses.append(loss)
    maes.append(mae)

# Resultados promedios
print(f"R2 promedio: {np.mean(r2_scores):.4f} ± {np.std(r2_scores):.4f}")
print(f"MSE promedio: {np.mean(losses):.4f} ± {np.std(losses):.4f}")
print(f"MAE promedio: {np.mean(maes):.4f} ± {np.std(maes):.4f}")


# Grafica de los CV K folds
r2_scores = np.array(r2_scores)
losses = np.array(losses)
maes = np.array(maes)


fig, axs = plt.subplots(3, 1, figsize=(10, 12))

# R2
axs[0].errorbar(range(1, 11), r2_scores, yerr=np.std(r2_scores), fmt='o-', color='#FF7F50', label='R² por fold')
axs[0].axhline(np.mean(r2_scores), color='b', linestyle='--', label=f'R² promedio: {np.mean(r2_scores):.4f}')
axs[0].set_title('R² por fold')
axs[0].set_xlabel('Fold')
axs[0].set_ylabel('R²')
axs[0].legend()
axs[0].grid(True)

# MSE
axs[1].errorbar(range(1, 11), losses, yerr=np.std(losses), fmt='o-', color='skyblue', label='Loss por fold')
axs[1].axhline(np.mean(losses), color='g', linestyle='--', label=f'Loss promedio: {np.mean(losses):.4f}')
axs[1].set_title('MSE (Loss) por fold')
axs[1].set_xlabel('Fold')
axs[1].set_ylabel('MSE')
axs[1].legend()
axs[1].grid(True)

# MAE
axs[2].errorbar(range(1, 11), maes, yerr=np.std(maes), fmt='o-', color='lightgreen', label='MAE por fold')
axs[2].axhline(np.mean(maes), color='m', linestyle='--', label=f'MAE promedio: {np.mean(maes):.4f}')
axs[2].set_title('MAE por fold')
axs[2].set_xlabel('Fold')
axs[2].set_ylabel('MAE')
axs[2].legend()
axs[2].grid(True)
plt.savefig("fold_separado.png", dpi=600, bbox_inches='tight')
plt.tight_layout()
plt.show()


##### Predicción número de especies para Coleoptera de Colombia con el mejor modelo por departamento y familia #####

'''
Cambiar el modelo según el que tenga el mejor resultado entre:

model_linear
dnn_simple
best_model_dnn_escalado
best_model_dnn
best_model_rf
'''

model = best_model_dnn_escalado
 
# Quitar la columna NumeroEspeciesFamilia del conjunto a predecir
coleoptera_completa_colombia = coleoptera_completa_colombia.loc[:, coleoptera_completa_colombia.columns != "NumeroEspeciesFamilia"] 
# Escalar los datos en caso que use el modelo escalado
scaler = StandardScaler()
scaled_array = scaler.fit_transform(coleoptera_completa_colombia)


prediccion_especies = model.predict(coleoptera_completa_colombia)

# Unir las predicciones a la tabla
estimacion_colombia = pd.DataFrame(coleoptera_completa_colombia) 
estimacion_colombia['prediccion_especies'] = prediccion_especies 

for column, encoder in label_encoders.items():
    if column in estimacion_colombia.columns:
        estimacion_colombia[column] = encoder.inverse_transform(estimacion_colombia[column])


# pivot table para la suma de especies únicas por familia en Colombia
tabla_resumen_estimaciones = pd.pivot_table(estimacion_colombia, values='prediccion_especies', index=['departamento','family'], aggfunc=['mean', 'median', 'max', 'min'])

# Redondear los datos al entero más cercan
pivot_table_rounded = np.ceil(tabla_resumen_estimaciones).astype(int)

# Exportar la tabla dinamica
pivot_table_rounded.to_excel('pivot_table_predictions_colombia_deparmentos.xlsx')

###############################################################################


prediccion_especies_antioquia = model.predict(coleoptera_completa_colombia)


for column, encoder in label_encoders.items():
    if column in estimacion_colombia.columns:
        estimacion_colombia[column] = encoder.inverse_transform(estimacion_colombia[column])


# Unir las predicciones a la tabla
estimacion_colombia = pd.DataFrame(coleoptera_completa_colombia)  # Convert X_test to DataFrame if it's not already
estimacion_colombia['prediccion_especies_antioquia'] = prediccion_especies_antioquia  # Add the predicted values as a new column


# pivot table para la suma de especies únicas por familia en Colombia
tabla_resumen_estimaciones = pd.pivot_table(estimacion_colombia, values='prediccion_especies_antioquia', index=['family'], aggfunc=['mean', 'median', 'max', 'min'])

# Redondear los datos al entero más cercan
pivot_table_rounded = np.ceil(tabla_resumen_estimaciones).astype(int)

# Exportar la tabla dinamica
pivot_table_rounded.to_excel('tabla_resumen_estimaciones_colombia.xlsx')



##### Predicción número de especies para Coleoptera de Antioquia con el mejor modelo

# Cargar solamente los datos ajustados
# coleoptera_completa_colombia_sinna = pd.read_csv('coleoptera_completa_sinna.txt', encoding = "utf8", sep="\t")
# coleopteros_colombia_etiquetados_antioquia = coleopteros_colombia_etiquetados_antioquia.drop(columns=['Coordinates','gbifID','occurrenceID','kingdom','phylum','class', 'order', 'stateProvince','infraspecificEpithet','license'])

coleopteros_colombia_etiquetados_antioquia=coleoptera_completa_colombia.loc[coleoptera_completa_colombia['departamento'] == 'Antioquia' ]



categorical_columns = coleopteros_colombia_etiquetados_antioquia.select_dtypes(include=['object']).columns
label_encoders = {}

for column in categorical_columns:
    le = LabelEncoder()  # Crear un LabelEncoder para cada columnas
    coleopteros_colombia_etiquetados_antioquia[column] = le.fit_transform(coleopteros_colombia_etiquetados_antioquia[column])
    label_encoders[column] = le # Guardar los diccionarios para poder hacer la trasnformación inversa luego


# Reconstruct the DataFrame with original column names and index
coleopteros_colombia_etiquetados_antioquia = pd.DataFrame(
    scaled_array,
    columns=coleopteros_colombia_etiquetados_antioquia.columns,
    index=coleopteros_colombia_etiquetados_antioquia.index
)

prediccion_especies_antioquia = model.predict(coleopteros_colombia_etiquetados_antioquia)    
    
# Unir las predicciones a la tabla
estimacion_antioquia = pd.DataFrame(coleopteros_colombia_etiquetados_antioquia) 
estimacion_antioquia['prediccion_especies'] = prediccion_especies_antioquia 

for column, encoder in label_encoders.items():
    estimacion_antioquia[column] = encoder.inverse_transform(estimacion_antioquia[column])


# pivot table para la suma de especies únicas por familia en Colombia
tabla_resumen_estimaciones_antioquia = pd.pivot_table(
    estimacion_antioquia,
    values='prediccion_especies',
    index=['family'],
    aggfunc=['mean', 'median', 'max', 'min']
)
# Redondear los datos al entero más cercan
pivot_table_rounded = np.ceil(tabla_resumen_estimaciones_antioquia).astype(int)

# Exportar la tabla dinamica
pivot_table_rounded.to_excel('pivot_table_predicciones_Antioquia.xlsx')



##### Grafica para mostrar el número de especies por familia en Antioquia
'''
Antes de correr esta parte, es necesario limpiar el resultado del paso anterior
para eliminar las subfamilias incorrectas y luego guardar un archivo .txt
con la información que se desea graficar, en general son dos columnas
la primera para las familias y la segunda con los valores predichos
'''

datos_prediccion_familia_antioquia = pd.read_csv('datosAntioquia.txt', encoding = "utf8", sep="\t")
# Ordenar datos para visualización
datos_ordenados = datos_prediccion_familia_antioquia.sort_values(by='Promedio de especies estimado Antioquia', ascending=True)

familias = datos_ordenados['Familia']
especies = datos_ordenados['Promedio de especies estimado Antioquia']

# Crear gráfico de barras horizontal
plt.figure(figsize=(10, 8))
bars = plt.barh(familias, especies, color='lightgreen')

# Añadir etiquetas del número de especies al final de cada barra
for bar in bars:
    ancho = bar.get_width()
    plt.text(ancho + 1,              
             bar.get_y() + bar.get_height() / 2, 
             f'{ancho:.0f}', 
             va='center', fontsize=9, color='black')

# Etiquetas y título
plt.xlabel('Número de especies')
plt.ylabel('Familia')
plt.title('Estimado del número de especies por familia en Antioquia')
plt.grid(axis='x')
plt.tight_layout()

# Calcular el total de la columna 'Promedio de especies estimado Antioquia'
total_especies = especies.sum()
plt.figtext(0.99, 0.01, f'Total de especies: {total_especies:.0f}', ha='right', va='bottom', fontsize=10, color='black')

# Guardar y mostrar
plt.savefig("barras_familia_antioquia.png", dpi=600, bbox_inches='tight')
plt.show()
